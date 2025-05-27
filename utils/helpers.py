import os
import sys
import json
import logging
import re
import tempfile
import time
import requests
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
import unicodedata
import chardet
from datetime import datetime
from pathlib import Path
import shutil
import hashlib 
from typing import Dict, Any, Optional, List, Union, Tuple

# Conditional imports for optional dependencies (file converters)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False
    print("Warning: python-docx library not found. DOCX file conversion will not be available.")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False
    print("Warning: PyPDF2 library not found. PDF file conversion will not be available.")

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    yaml = None
    YAML_AVAILABLE = False
    print("Warning: PyYAML library not found. YAML configuration loading will use defaults only.")

module_logger = logging.getLogger(__name__)
if not module_logger.hasHandlers():
    module_logger.addHandler(logging.NullHandler())


def setup_logger(log_dir: str = 'logs',
                 log_name: str = 'SyllabusPipelineHelper',
                 log_level: int = logging.INFO,
                 config_format: Optional[str] = None,
                 config_datefmt: Optional[str] = None
                ) -> logging.Logger:
    log_dir_path = Path(log_dir)
    try:
        log_dir_path.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        print(f"Warning: Could not create log directory {log_dir_path}: {e}. Using temp dir for logs.", file=sys.stderr)
        log_dir_path = Path(tempfile.gettempdir()) / "syllabus_parser_logs"
        log_dir_path.mkdir(parents=True, exist_ok=True)

    local_logger = logging.getLogger(log_name)
    if local_logger.hasHandlers():
        local_logger.handlers.clear()
    local_logger.setLevel(log_level)
    local_logger.propagate = False
    timestamp = datetime.now().strftime('%Y%m%d')
    log_file = log_dir_path / f'{log_name}_{timestamp}.log'
    try:
        file_handler = logging.FileHandler(log_file, mode='a', encoding='utf-8')
        file_handler.setLevel(log_level)
        console_handler = logging.StreamHandler(sys.stdout)
        console_handler.setLevel(log_level)
        log_format_str = config_format or '%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s'
        date_format_str = config_datefmt or '%Y-%m-%d %H:%M:%S'
        formatter = logging.Formatter(log_format_str, datefmt=date_format_str)
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        local_logger.addHandler(file_handler)
        local_logger.addHandler(console_handler)
        local_logger.info(f"Logger '{log_name}' configured. Level: {logging.getLevelName(log_level)}. File: {log_file}")
    except Exception as e_log_setup:
        print(f"CRITICAL: Failed to set up logger '{log_name}': {e_log_setup}", file=sys.stderr)
        logging.basicConfig(level=logging.WARNING, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        local_logger = logging.getLogger(log_name)
    return local_logger

def load_configuration(config_path: Optional[Union[str, Path]] = None,
                       default_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    if default_config is None:
        project_root_default = Path(__file__).resolve().parent.parent
        default_config = {
            'project_root': str(project_root_default),
            'directories': {
                'uploads': str(project_root_default / 'uploads'),
                'converted': str(project_root_default / 'converted'),
                'output': str(project_root_default / 'output'),
                'logs': str(project_root_default / 'logs'),
                'original_syllabus_dir': str(project_root_default / 'syllabus_data' / 'originals'),
                'converted_syllabus_dir': str(project_root_default / 'syllabus_data' / 'converted_text'),
                'parsed_syllabus_dir': str(project_root_default / 'syllabus_data' / 'parsed_results'),
                'main_data_dir': str(project_root_default / 'syllabus_data' / 'main_data_store')
            },
            'file_settings': {
                'allowed_extensions': ['pdf', 'docx', 'txt', 'html', 'htm'],
                'max_content_length': 16 * 1024 * 1024
            },
            'flask': {
                'secret_key': os.environ.get('FLASK_SECRET_KEY', 'default_flask_secret_key_please_change_in_production'),
                'port': int(os.environ.get('FLASK_PORT', 5000)),
                'host': os.environ.get('FLASK_HOST', '0.0.0.0'),
                'debug': os.environ.get('FLASK_DEBUG', 'True').lower() == 'true'
            },
            'logging': {
                'level': 'INFO',
                'format': '%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s',
                'datefmt': '%Y-%m-%d %H:%M:%S'
            },
            'extraction': {'openai_model': 'gpt-4o', 'openai_api_key': None},
            'openai_parser': {'max_chars_for_metadata': 30000, 'max_api_retries': 2},
            'syllabus_extractor': {'initial_snippet_max_chars': 1000}, 
            'schedule_parser': {'default_term_weeks': 15},
            'required_fields': ["Course Title", "Term", "Days of Week", "Class Start Date", "Class End Date", "Instructor Name", "Time Zone"]
        }
    config = default_config.copy()
    if config_path:
        config_file = Path(config_path)
        if YAML_AVAILABLE and config_file.exists() and config_file.is_file():
            try:
                with config_file.open("r", encoding="utf-8") as f:
                    yaml_config = yaml.safe_load(f)
                if yaml_config and isinstance(yaml_config, dict):
                    for section, values in yaml_config.items():
                        if section in config and isinstance(config[section], dict) and isinstance(values, dict):
                            config[section].update(values)
                        else:
                            config[section] = values
                    module_logger.info(f"Configuration successfully loaded and merged from {config_file}")
                elif not yaml_config:
                     module_logger.warning(f"Configuration file {config_file} is empty. Using defaults.")
                else:
                     module_logger.warning(f"Configuration file {config_file} did not load as a dictionary. Using defaults.")
            except yaml.YAMLError as e_yaml:
                module_logger.error(f"Error parsing YAML configuration file {config_file}: {e_yaml}", exc_info=True)
            except Exception as e_load:
                module_logger.error(f"Error loading configuration file {config_file}: {e_load}", exc_info=True)
        elif not YAML_AVAILABLE and config_file.exists():
            module_logger.warning(f"YAML library not available, but config file {config_file} exists. Using default configuration.")
        elif not config_file.exists():
            module_logger.warning(f"Configuration file not found: {config_file}. Using default configuration.")
    else:
        module_logger.info("No configuration file path provided. Using default configuration.")
    return config

def setup_environment(config: Dict[str, Any], project_root: Optional[Union[str,Path]] = None) -> bool:
    try:
        if project_root is None:
            project_root_path = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
        else:
            project_root_path = Path(project_root)
        if str(project_root_path) not in sys.path:
            sys.path.insert(0, str(project_root_path))
            module_logger.debug(f"Added project root to sys.path: {project_root_path}")
        utils_path = project_root_path / "utils"
        parsers_path = utils_path / "parsers"
        if not (utils_path / "__init__.py").exists():
            module_logger.warning(f"Missing __init__.py in {utils_path}.")
        if not (parsers_path / "__init__.py").exists():
            module_logger.warning(f"Missing __init__.py in {parsers_path}.")
        ensure_directories(config, project_root_path)
        module_logger.info("Environment setup completed.")
        return True
    except Exception as e:
        module_logger.error(f"Error setting up environment: {e}", exc_info=True)
        return False

def ensure_directories(config: Dict[str, Any], project_root_path: Path) -> bool:
    try:
        directories_config = config.get('directories', {})
        if not directories_config:
            module_logger.warning("No 'directories' section in config to ensure.")
            return True
        for dir_key, dir_path_str in directories_config.items():
            if not isinstance(dir_path_str, str):
                module_logger.warning(f"Directory path for key '{dir_key}' is not a string: {dir_path_str}. Skipping.")
                continue
            dir_path = Path(dir_path_str)
            if not dir_path.is_absolute():
                cfg_project_root = Path(config.get('project_root', project_root_path))
                dir_path = cfg_project_root / dir_path
            dir_path.mkdir(parents=True, exist_ok=True)
            module_logger.debug(f"Ensured directory exists: {dir_path} (for config key: {dir_key})")
        return True
    except Exception as e:
        module_logger.error(f"Error ensuring directories: {e}", exc_info=True)
        return False

def allowed_file(filename: str, allowed_extensions: Optional[List[str]] = None) -> bool:
    if allowed_extensions is None:
        allowed_extensions = ['pdf', 'docx', 'txt', 'html', 'htm']
    if not filename or not isinstance(filename, str): return False
    secure_name = secure_filename(filename)
    return '.' in secure_name and secure_name.rsplit('.', 1)[1].lower() in set(ext.lower() for ext in allowed_extensions)

def read_json(file_path: Union[str, Path], fallback_data: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    path_obj = Path(file_path)
    effective_fallback = fallback_data if fallback_data is not None else {}
    if not path_obj.is_file():
        module_logger.warning(f"JSON file does not exist: {path_obj}")
        return effective_fallback
    try:
        with path_obj.open('rb') as rawfile:
            raw_data = rawfile.read()
            if not raw_data:
                module_logger.warning(f"JSON file is empty: {path_obj}")
                return effective_fallback
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] if detected['encoding'] else 'utf-8'
        with path_obj.open('r', encoding=encoding, errors='replace') as file:
            data = json.load(file)
        return data
    except json.JSONDecodeError as e_decode:
        module_logger.warning(f"JSON decode error in {path_obj} (encoding {encoding}): {e_decode}. Attempting repair.")
        try:
            with path_obj.open('r', encoding=encoding, errors='replace') as file: content = file.read()
            content_repaired = re.sub(r',\s*([\}\]])', r'\1', content)
            if content_repaired != content: module_logger.info(f"Attempted JSON repair for {path_obj}.")
            data = json.loads(content_repaired)
            return data
        except Exception as e_repair:
            module_logger.error(f"Failed to repair/parse JSON from {path_obj}: {e_repair}", exc_info=True)
            return effective_fallback
    except Exception as e_read:
        module_logger.error(f"Error reading JSON {path_obj}: {e_read}", exc_info=True)
        return effective_fallback

def write_json(file_path: Union[str, Path], data: Dict[str, Any], create_dirs: bool = True, backup_existing: bool = True, indent: int = 4) -> bool:
    path_obj = Path(file_path)
    try:
        if create_dirs: path_obj.parent.mkdir(parents=True, exist_ok=True)
        if backup_existing and path_obj.exists() and path_obj.is_file():
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S%f')
            backup_path = path_obj.with_name(f"{path_obj.stem}_{timestamp}{path_obj.suffix}.bak")
            try:
                shutil.copy2(str(path_obj), str(backup_path))
            except Exception as backup_err:
                module_logger.warning(f"Failed to backup {path_obj}: {backup_err}")
        with path_obj.open('w', encoding='utf-8') as file:
            json.dump(data, file, indent=indent, ensure_ascii=False, default=str)
        try: 
            with path_obj.open('r', encoding='utf-8') as check_file: json.load(check_file)
            return True
        except Exception as verify_err:
            module_logger.error(f"File {path_obj} written but failed verification: {verify_err}")
            return False
    except Exception as e:
        module_logger.error(f"Error writing JSON to {path_obj}: {e}", exc_info=True)
        return False

def normalize_text(text: Optional[str]) -> str:
    if text is None: return ""
    try:
        text_normalized = unicodedata.normalize('NFC', str(text))
        text_no_exotic_spaces = re.sub(r'[\s\u00A0]+', ' ', text_normalized)
        return text_no_exotic_spaces.strip()
    except Exception as e:
        module_logger.warning(f"Error normalizing text (returning original stripped): {e}")
        return str(text).strip()

def generate_content_hash(content: Union[bytes, str], hash_algo: str = 'sha256') -> str:
    hasher = hashlib.new(hash_algo)
    if isinstance(content, str):
        content_bytes = content.encode('utf-8')
    elif isinstance(content, bytes):
        content_bytes = content
    else:
        raise TypeError("Content must be bytes or string to hash.")
    hasher.update(content_bytes)
    return hasher.hexdigest()

def _read_and_normalize_text_file_content(file_path: Path, detected_encoding: str) -> str:
    with file_path.open('r', encoding=detected_encoding, errors='replace') as f: text = f.read()
    return normalize_text(text)

def convert_txt(file_path: Path, output_path: Optional[Path] = None) -> str:
    # ... (Keep implementation as is) ...
    module_logger.info(f"Processing TXT file: {file_path}")
    try:
        with file_path.open('rb') as rawfile: raw_data = rawfile.read()
        if not raw_data: module_logger.warning(f"TXT file is empty: {file_path}"); return ""
        detected = chardet.detect(raw_data)
        encoding_to_try = detected['encoding'] if detected['encoding'] and detected.get('confidence', 0) > 0.5 else 'utf-8'
        try: text = _read_and_normalize_text_file_content(file_path, encoding_to_try)
        except UnicodeDecodeError:
            module_logger.warning(f"TXT: Failed with {encoding_to_try}. Trying fallbacks for {file_path.name}.")
            common_encodings = ['utf-8', 'latin-1', 'windows-1252']
            if encoding_to_try in common_encodings: common_encodings.remove(encoding_to_try)
            text_content_found = False
            for enc in common_encodings:
                try:
                    text = _read_and_normalize_text_file_content(file_path, enc)
                    text_content_found = True; break
                except UnicodeDecodeError: continue
            if not text_content_found: module_logger.error(f"TXT: Could not decode {file_path.name}."); return ""
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(text, encoding='utf-8')
        return text
    except Exception as e: module_logger.error(f"Error processing TXT {file_path}: {e}", exc_info=True); return ""

def convert_pdf_to_txt(file_path: Path, output_path: Optional[Path] = None) -> str:
    # ... (Keep implementation as is) ...
    if not PYPDF2_AVAILABLE: module_logger.error("PyPDF2 not available for PDF conversion."); return ""
    module_logger.info(f"Processing PDF file: {file_path}")
    try:
        text_parts = []
        with file_path.open('rb') as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                module_logger.warning(f"PDF {file_path.name} encrypted. Attempting decrypt.")
                try:
                    if hasattr(PyPDF2, 'PasswordType'): 
                        if reader.decrypt('') == PyPDF2.PasswordType.DECRYPTION_FAILED: module_logger.warning(f"PDF {file_path.name} decrypt failed (v3+).")
                    elif not reader.decrypt(''): module_logger.warning(f"PDF {file_path.name} decrypt failed (older).")
                except Exception as decrypt_err: module_logger.warning(f"PDF {file_path.name} decrypt error: {decrypt_err}")
            for i, page in enumerate(reader.pages):
                try:
                    extracted_text = page.extract_text()
                    if extracted_text: text_parts.append(extracted_text)
                except Exception as page_err: module_logger.warning(f"Error PDF page {i+1} of {file_path.name}: {page_err}")
        normalized_text = normalize_text("\n\n".join(text_parts))
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(normalized_text, encoding='utf-8')
        return normalized_text
    except Exception as e: module_logger.error(f"Error converting PDF {file_path.name}: {e}", exc_info=True); return ""

def convert_docx_to_txt(file_path: Path, output_path: Optional[Path] = None) -> str:
    # ... (Keep implementation as is) ...
    if not DOCX_AVAILABLE or DocxDocument is None: module_logger.error("python-docx not available for DOCX conversion."); return ""
    module_logger.info(f"Processing DOCX file: {file_path}")
    try:
        doc = DocxDocument(str(file_path))
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            table_text_parts = [" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()) for row in table.rows]
            if any(table_text_parts): text_parts.append("\n".join(filter(None, table_text_parts)))
        normalized_text = normalize_text("\n\n".join(text_parts))
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(normalized_text, encoding='utf-8')
        return normalized_text
    except Exception as e: module_logger.error(f"Error converting DOCX {file_path.name}: {e}", exc_info=True); return ""

def convert_webpage_to_txt(url: str, output_path: Optional[Union[str, Path]] = None) -> Optional[str]:
    # ... (Keep implementation as is) ...
    module_logger.info(f"Fetching webpage: {url}")
    try:
        if not url.startswith(('http://', 'https://')): url = 'https://' + url
        headers = {'User-Agent': 'Mozilla/5.0', 'Accept': 'text/html', 'Accept-Language': 'en-US,en;q=0.5'}
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        content_type = response.headers.get('Content-Type', '').lower()
        if 'application/pdf' in content_type:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                temp_pdf.write(response.content); temp_pdf_path = Path(temp_pdf.name)
            try: return convert_pdf_to_txt(temp_pdf_path, Path(output_path) if output_path else None)
            finally:
                if temp_pdf_path.exists(): temp_pdf_path.unlink()
        if 'text/html' not in content_type: return normalize_text(response.text)
        soup = BeautifulSoup(response.content, 'html.parser')
        page_title = soup.title.string.strip() if soup.title and soup.title.string else "Untitled"
        for element_type in soup(['script', 'style', 'head', 'header', 'footer', 'nav', 'aside', 'form', 'noscript', 'iframe', 'link', 'meta']):
            element_type.decompose()
        main_content_selectors = ['main', 'article', '[role="main"]', '.content', '.main-content', '#content', '#main', '.entry-content']
        content_element = next((soup.select_one(s) for s in main_content_selectors if soup.select_one(s)), soup.body or soup)
        text_blocks = [element.strip() for element in content_element.find_all(string=True) if element.strip()]
        normalized_text = normalize_text("\n".join(text_blocks))
        contextual_text = f"Title: {page_title}\nURL: {url}\n\n{normalized_text}"
        if output_path:
            out_path = Path(output_path); out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_text(contextual_text, encoding='utf-8')
        return contextual_text
    except requests.exceptions.RequestException as e_req: module_logger.error(f"Request error for {url}: {e_req}", exc_info=True)
    except Exception as e_web: module_logger.error(f"Error converting webpage {url}: {e_web}", exc_info=True)
    return None

def convert_file_to_text(file_path: Optional[Union[str, Path]] = None,
                         file_content: Optional[bytes] = None,
                         file_extension: Optional[str] = None,
                         output_path: Optional[Union[str, Path]] = None) -> Optional[str]:
    # ... (Keep implementation as is) ...
    output_path_obj = Path(output_path) if output_path else None
    try:
        if file_path:
            path_obj = Path(file_path)
            if not path_obj.is_file(): module_logger.error(f"File not found: {path_obj}"); return None
            ext = path_obj.suffix.lower().lstrip(".")
            if ext == 'pdf': return convert_pdf_to_txt(path_obj, output_path_obj)
            elif ext == 'docx': return convert_docx_to_txt(path_obj, output_path_obj)
            elif ext == 'txt': return convert_txt(path_obj, output_path_obj)
            elif ext in ['html', 'htm']:
                html_content = path_obj.read_bytes().decode(chardet.detect(path_obj.read_bytes())['encoding'] or 'utf-8', errors='replace')
                soup = BeautifulSoup(html_content, 'html.parser')
                for el_type in soup(['script','style','head','header','footer','nav','aside']): el_type.decompose()
                text = normalize_text(soup.get_text(separator='\n'))
                if output_path_obj: output_path_obj.parent.mkdir(parents=True, exist_ok=True); output_path_obj.write_text(text, encoding='utf-8')
                return text
            else: module_logger.error(f"Unsupported extension '{ext}' from path: {path_obj}"); return None
        elif file_content and file_extension:
            ext = file_extension.lower().lstrip(".")
            with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as temp_file:
                temp_file.write(file_content); temp_file_path = Path(temp_file.name)
            text_result = None
            try:
                if ext == 'pdf': text_result = convert_pdf_to_txt(temp_file_path, output_path_obj)
                elif ext == 'docx': text_result = convert_docx_to_txt(temp_file_path, output_path_obj)
                elif ext == 'txt': text_result = convert_txt(temp_file_path, output_path_obj)
                elif ext in ['html', 'htm']:
                    html_content = temp_file_path.read_bytes().decode(chardet.detect(temp_file_path.read_bytes())['encoding'] or 'utf-8', errors='replace')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    for el_type in soup(['script','style','head','header','footer','nav','aside']): el_type.decompose()
                    text_result = normalize_text(soup.get_text(separator='\n'))
                    if output_path_obj: output_path_obj.parent.mkdir(parents=True, exist_ok=True); output_path_obj.write_text(text_result, encoding='utf-8')
                else: module_logger.error(f"Unsupported extension for content: '{ext}'")
            finally:
                if temp_file_path.exists(): temp_file_path.unlink()
            return text_result
        else: module_logger.error("File path or (content+extension) required for convert_file_to_text."); return None
    except Exception as e: module_logger.error(f"Error in convert_file_to_text: {e}", exc_info=True); return None

def save_uploaded_file(file_obj: Any, upload_folder: Path, unique_id_prefix: str) -> Tuple[Optional[Path], Optional[str]]:
    # ... (Keep implementation as is) ...
    try:
        filename_attr = getattr(file_obj, 'filename', '')
        if not filename_attr: filename = f"upload_{unique_id_prefix.split('-')[0]}.dat"
        else: filename = secure_filename(filename_attr)
        if not filename: filename = f"upload_secure_{unique_id_prefix.split('-')[0]}.dat"
        original_extension = Path(filename).suffix
        final_filename_in_uploads = f"{unique_id_prefix}{original_extension}" if original_extension else f"{unique_id_prefix}.dat"
        upload_folder.mkdir(parents=True, exist_ok=True)
        file_path = upload_folder / final_filename_in_uploads
        file_obj.save(str(file_path))
        module_logger.info(f"Saved uploaded file to: {file_path}")
        return file_path, final_filename_in_uploads 
    except Exception as e:
        module_logger.error(f"Error saving uploaded file '{getattr(file_obj, 'filename', 'N/A')}': {e}", exc_info=True)
        return None, None

def extract_course_code(text: str, config: Optional[Dict[str, Any]] = None) -> Optional[str]:
    # ... (Keep implementation as is) ...
    if not text or not isinstance(text, str): return None
    patterns = [
        r'\b([A-Z]{2,4}\s+\d{3,4}[A-Z]?)(?:\.\d+)?\b',
        r'\b([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\b',
        r'\b([A-Z]{2,4}\d{2,3}[A-Z]?)\b'
    ]
    for pattern_str in patterns:
        try:
            match = re.search(pattern_str, text, re.IGNORECASE)
            if match:
                course_code_raw = match.group(1)
                return re.sub(r'[-\s]', '', course_code_raw).upper()
        except re.error as e_re: module_logger.warning(f"Regex error in extract_course_code: {e_re}"); continue
    return None


def process_syllabus_pipeline(
    input_path_or_text: Union[str, Path],
    config: Dict[str, Any],
    parsers: Dict[str, Any],
    original_filename_for_display: Optional[str] = None, 
    transient_run_id: Optional[str] = None
) -> Dict[str, Any]:
    pipeline_start_time = time.monotonic()
    module_logger.info(f"PIPELINE START: Input: '{str(input_path_or_text)[:100]}...', Original Filename: {original_filename_for_display}, Transient ID: {transient_run_id}")

    # --- Hashing, Path Setup, Caching (keep as is, with caching commented out) ---
    # ... (This part should be the same as helpers_py_task_data_fix_v3) ...
    content_hash: Optional[str] = None
    original_file_for_conversion: Optional[Path] = None
    original_file_extension: str = ".dat" 
    if isinstance(input_path_or_text, (str, Path)) and Path(input_path_or_text).is_file():
        original_file_for_conversion = Path(input_path_or_text)
        original_file_extension = original_file_for_conversion.suffix
        try:
            with original_file_for_conversion.open('rb') as f: content_hash = generate_content_hash(f.read())
            module_logger.info(f"Generated hash for file {original_file_for_conversion.name}: {content_hash}")
        except Exception as e_hash:
            module_logger.error(f"Error generating hash for file {original_file_for_conversion.name}: {e_hash}", exc_info=True)
            content_hash = transient_run_id or str(datetime.now().timestamp())
            module_logger.warning(f"Using fallback ID {content_hash} as primary identifier due to hashing error.")
    elif isinstance(input_path_or_text, str): 
        raw_text_content = input_path_or_text
        content_hash = generate_content_hash(raw_text_content)
        original_file_extension = ".txt" 
        module_logger.info(f"Generated hash for raw text input: {content_hash}")
        try:
            temp_dir_path_str = config.get("directories", {}).get("uploads", "temp_uploads_for_text")
            cfg_project_root_temp = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
            temp_dir = Path(temp_dir_path_str)
            if not temp_dir.is_absolute(): temp_dir = cfg_project_root_temp / temp_dir
            temp_dir.mkdir(parents=True, exist_ok=True)
            original_file_for_conversion = temp_dir / f"{content_hash}_raw_text_input.txt"
            original_file_for_conversion.write_text(raw_text_content, encoding='utf-8')
            module_logger.info(f"Raw text input saved to temporary file: {original_file_for_conversion}")
        except Exception as e_temp_save:
            module_logger.error(f"Error saving raw text input to temporary file: {e_temp_save}", exc_info=True)
    else:
        module_logger.error(f"Invalid input type for pipeline: {type(input_path_or_text)}")
        return {"metadata": {"error": "Invalid input type", "unique_id": transient_run_id or "error_id"}}
    if not content_hash: 
        content_hash = transient_run_id or str(datetime.now().timestamp())
        module_logger.error(f"Content hash is unexpectedly None after input processing. Using fallback ID: {content_hash}")
    
    module_logger.info(f"CACHE MISS or CACHING DISABLED: Processing anew for hash {content_hash}.")

    results: Dict[str, Any] = {
        "metadata": {
            "pipeline_start_time": datetime.now().isoformat(), "unique_id": content_hash, 
            "transient_run_id": transient_run_id, 
            "original_file": original_filename_for_display or (original_file_for_conversion.name if original_file_for_conversion else "text_input"),
            "extraction_complete": False, "parser_stages_completed": [], "text_length": 0,
            "Scheduled Labs": "No", "Recitation Sessions": "No"
        },
        "class_data": {}, "event_data": [], "task_data": [], 
        "lab_data": [], "recitation_data": [], "segmented_syllabus": {}
    }

    # --- Storing Original File (keep as is) ---
    # ... (This part should be the same as helpers_py_task_data_fix_v3) ...
    if original_file_for_conversion and original_file_for_conversion.is_file(): 
        original_syllabus_storage_dir_str = config.get("directories", {}).get("original_syllabus_dir", "syllabus_data/originals")
        cfg_project_root_orig = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
        original_syllabus_storage_dir = Path(original_syllabus_storage_dir_str)
        if not original_syllabus_storage_dir.is_absolute(): original_syllabus_storage_dir = cfg_project_root_orig / original_syllabus_storage_dir
        original_syllabus_storage_dir.mkdir(parents=True, exist_ok=True)
        stored_original_path = original_syllabus_storage_dir / f"{content_hash}{original_file_extension}"
        try:
            if original_file_for_conversion and isinstance(input_path_or_text, (str, Path)) and Path(input_path_or_text) == original_file_for_conversion: 
                 shutil.copy2(original_file_for_conversion, stored_original_path)
                 module_logger.info(f"Stored original uploaded file at: {stored_original_path}")
            elif original_file_for_conversion and isinstance(input_path_or_text, str): 
                 shutil.copy2(original_file_for_conversion, stored_original_path)
                 module_logger.info(f"Stored original (from raw text input temp file) at: {stored_original_path}")
        except Exception as e_store_orig: module_logger.warning(f"Could not store original file copy at {stored_original_path}: {e_store_orig}")

    # --- Text Conversion (keep as is) ---
    # ... (This part should be the same as helpers_py_task_data_fix_v3, ensure converted_text_output_path_persistent is handled) ...
    converted_text_dir_str = config.get("directories", {}).get("converted_syllabus_dir", "syllabus_data/converted_text")
    cfg_project_root_conv = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
    converted_text_dir = Path(converted_text_dir_str)
    if not converted_text_dir.is_absolute(): converted_text_dir = cfg_project_root_conv / converted_text_dir
    converted_text_dir.mkdir(parents=True, exist_ok=True)
    converted_text_output_path_persistent = converted_text_dir / f"{content_hash}_converted.txt"
    syllabus_text_content: Optional[str] = None
    if original_file_for_conversion and original_file_for_conversion.is_file():
        syllabus_text_content = convert_file_to_text(file_path=original_file_for_conversion, output_path=converted_text_output_path_persistent)
    elif isinstance(input_path_or_text, str): 
        syllabus_text_content = input_path_or_text 
        try: converted_text_output_path_persistent.write_text(syllabus_text_content, encoding='utf-8'); module_logger.info(f"Raw text directly saved to persistent converted location: {converted_text_output_path_persistent}")
        except Exception as e_raw_save: module_logger.error(f"Failed to save raw text to persistent converted location {converted_text_output_path_persistent}: {e_raw_save}")
    if not syllabus_text_content:
        module_logger.error(f"PIPELINE ID {content_hash}: Failed to convert or obtain input text. Aborting.")
        results["metadata"]["error"] = "Failed to convert or obtain input text."; results["metadata"]["process_time_pipeline"] = round(time.monotonic() - pipeline_start_time, 3)
        return results 
    results["metadata"]["text_length"] = len(syllabus_text_content)
    if not converted_text_output_path_persistent.exists() and syllabus_text_content:
        module_logger.warning(f"Persistent converted text file {converted_text_output_path_persistent} did not exist. Writing now.")
        try: converted_text_output_path_persistent.write_text(syllabus_text_content, encoding='utf-8')
        except Exception as e_write_late: module_logger.error(f"Failed to write to {converted_text_output_path_persistent} for SyllabusParser: {e_write_late}")


    # --- SyllabusParser (Segmentation) ---
    # ... (Keep as is) ...
    syllabus_parser_instance = parsers.get('syllabus_parser')
    if syllabus_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 1 - Running SyllabusParser using {converted_text_output_path_persistent}.")
        try:
            segmentation_output = syllabus_parser_instance.parse_syllabus(input_text_file_path=str(converted_text_output_path_persistent), unique_id=content_hash)
            if segmentation_output.get("status") == "success" and segmentation_output.get("parsed_data"):
                results["segmented_syllabus"] = segmentation_output["parsed_data"]
                results["metadata"]["parser_stages_completed"].append("syllabus_segmentation")
            else:
                results["metadata"]["error_syllabus_parser"] = segmentation_output.get('error', "SyllabusParser failed without detailed error")
                module_logger.error(f"PIPELINE ID {content_hash}: SyllabusParser failed. Error: {results['metadata']['error_syllabus_parser']}")
        except Exception as e_sp: results["metadata"]["error_syllabus_parser"] = str(e_sp); module_logger.error(f"PIPELINE ID {content_hash}: Error during SyllabusParser stage: {e_sp}", exc_info=True)
    else: module_logger.warning(f"PIPELINE ID {content_hash}: SyllabusParser not available. Segmentation skipped.")

    # --- SyllabusExtractor (class_data, initial event_data, initial task_data) ---
    # ... (Keep as is) ...
    extractor_instance = parsers.get('extractor')
    if extractor_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 2 - Running Extractor (SyllabusExtractor).")
        try:
            extraction_output = extractor_instance.extract_all(full_syllabus_text=syllabus_text_content, segmented_data=results.get("segmented_syllabus", {}), unique_id=content_hash)
            if extraction_output.get("class_data"): results["class_data"].update(extraction_output["class_data"])
            results.setdefault("task_data", [])
            if "task_data" in extraction_output and isinstance(extraction_output["task_data"], list): results["task_data"].extend(extraction_output["task_data"]); module_logger.info(f"SyllabusExtractor added/extended {len(extraction_output['task_data'])} items to task_data.")
            results.setdefault("event_data", [])
            if "event_data" in extraction_output and isinstance(extraction_output["event_data"], list): results["event_data"].extend(extraction_output["event_data"]); module_logger.info(f"SyllabusExtractor added/extended {len(extraction_output['event_data'])} items to event_data.")
            if "metadata" in extraction_output: 
                results["metadata"]["extractor_model_used"] = extraction_output["metadata"].get("model_used")
                results["metadata"]["class_data_missing_fields_by_extractor"] = extraction_output["metadata"].get("missing_fields_by_parser") 
                results["metadata"]["extractor_process_time"] = extraction_output["metadata"].get("process_time_extractor")
                if extraction_output["metadata"].get("error"): results["metadata"]["error_extractor"] = extraction_output["metadata"]["error"]
            results["metadata"]["parser_stages_completed"].append("class_data_extraction_by_extractor")
        except Exception as e_ext: results["metadata"]["error_extractor"] = str(e_ext); module_logger.error(f"PIPELINE ID {content_hash}: Error during Extractor stage: {e_ext}", exc_info=True)
    else: module_logger.warning(f"PIPELINE ID {content_hash}: SyllabusExtractor not available. Primary class_data extraction skipped.")

    # --- DateParser (process class_data dates and other top-level date fields) ---
    # ... (Keep as is) ...
    date_parser_instance = parsers.get('date_parser')
    if date_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 3 - Running DateParser.process_dates.")
        try: 
            term_context_for_dp = results.get("class_data", {}).get("Term")
            results = date_parser_instance.process_dates(results, term_year_str=term_context_for_dp) 
            results["metadata"]["parser_stages_completed"].append("date_processing")
        except Exception as e_dp_proc: results["metadata"]["error_date_parser_processing"] = str(e_dp_proc); module_logger.error(f"DateParser process_dates error: {e_dp_proc}", exc_info=True)
    else: module_logger.warning(f"PIPELINE ID {content_hash}: DateParser not available.")
    
    # --- AssignmentParser (extract_tasks_from_text using specific segments) ---
    assignment_parser_instance = parsers.get('assignment_parser')
    if assignment_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 4 - Running AssignmentParser.extract_tasks_from_text.")
        try:
            assignment_text_segment = results.get("segmented_syllabus", {}).get("assignments_exams", "")
            if not isinstance(assignment_text_segment, str) or not assignment_text_segment.strip():
                assignment_text_segment = results.get("segmented_syllabus", {}).get("course_requirements_grading_policy", syllabus_text_content)
                if not isinstance(assignment_text_segment, str) or not assignment_text_segment.strip():
                    # Check for 'exam_and_homework_schedule' if SyllabusParser created it
                    assignment_text_segment = results.get("segmented_syllabus", {}).get("exam_and_homework_schedule", syllabus_text_content)
                    if isinstance(assignment_text_segment, str) and assignment_text_segment.strip() and assignment_text_segment != syllabus_text_content:
                         module_logger.info("Using 'exam_and_homework_schedule' segment for task extraction.")
                    elif isinstance(assignment_text_segment, str) and assignment_text_segment.strip() : # it was syllabus_text_content
                         module_logger.info(f"Using full syllabus text (len: {len(assignment_text_segment)}) for task extraction as other segments were empty.")
                    else: # It was an empty string from course_requirements_grading_policy
                        assignment_text_segment = syllabus_text_content
                        module_logger.info(f"Using ultimate fallback (full syllabus text, len: {len(assignment_text_segment)}) for task extraction.")

            if assignment_text_segment:
                newly_extracted_tasks_by_ap = assignment_parser_instance.extract_tasks_from_text(
                    assignment_text_segment, 
                    results.get("class_data", {})
                )
                if newly_extracted_tasks_by_ap and isinstance(newly_extracted_tasks_by_ap, list):
                    # Ensure results["task_data"] exists and is a list
                    if "task_data" not in results or not isinstance(results["task_data"], list):
                        results["task_data"] = [] 
                    
                    results["task_data"].extend(newly_extracted_tasks_by_ap) # Directly extend
                    module_logger.info(f"AssignmentParser.extract_tasks_from_text added {len(newly_extracted_tasks_by_ap)} tasks. Total task_data now: {len(results['task_data'])}.")
                elif newly_extracted_tasks_by_ap:
                    module_logger.error(f"AssignmentParser.extract_tasks_from_text returned tasks but not in list format: {type(newly_extracted_tasks_by_ap)}. Task data not updated.")
                else:
                    module_logger.info("AssignmentParser.extract_tasks_from_text returned no new tasks from its segment.")
            else:
                module_logger.warning("No suitable text segment found for assignment extraction by AssignmentParser.")
            
            results["metadata"]["parser_stages_completed"].append("assignment_extraction_llm")
        except Exception as e_ap_extract:
            results["metadata"]["error_assignment_parser_extraction"] = str(e_ap_extract) 
            module_logger.error(f"AssignmentParser extract_tasks_from_text error: {e_ap_extract}", exc_info=True)
    else:
        module_logger.warning(f"PIPELINE ID {content_hash}: AssignmentParser not available.")

    # --- LabParser (extract_labs_from_text) ---
    # ... (Keep LabParser logic as is from helpers_py_task_data_fix_v3) ...
    lab_parser_instance = parsers.get('lab_parser')
    if lab_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 5 - Running LabParser.extract_labs_from_text.")
        try:
            lab_text_segment = results.get("segmented_syllabus", {}).get("separate_laboratory_sections", "")
            if not isinstance(lab_text_segment, str) or not lab_text_segment.strip():
                lab_text_segment = results.get("segmented_syllabus", {}).get("phy273_lab_policies", "")
                if not isinstance(lab_text_segment, str) or not lab_text_segment.strip():
                    lab_text_segment = results.get("segmented_syllabus", {}).get("lab_schedule", "")
                    if not isinstance(lab_text_segment, str) or not lab_text_segment.strip():
                        lab_text_segment = syllabus_text_content 
                        module_logger.info(f"Using full text segment for lab extraction as specific lab sections were empty.")
                    else: module_logger.info("Using 'lab_schedule' segment for lab extraction.")
                else: module_logger.info("Using 'phy273_lab_policies' segment for lab extraction.")
            else: module_logger.info("Using 'separate_laboratory_sections' for lab extraction.")
            if lab_text_segment:
                newly_extracted_labs = lab_parser_instance.extract_labs_from_text(lab_text_segment, results.get("class_data", {}))
                if newly_extracted_labs:
                    results.setdefault("lab_data", []).extend(newly_extracted_labs) 
                    results["metadata"]["Scheduled Labs"] = "Yes"
                    module_logger.info(f"LabParser.extract_labs_from_text added {len(newly_extracted_labs)} labs. Total lab_data: {len(results['lab_data'])}.")
                else: module_logger.info("LabParser.extract_labs_from_text returned no new labs.")
            results["metadata"]["parser_stages_completed"].append("lab_data_extraction")
        except Exception as e_lp_extract: results["metadata"]["error_lab_parser_extraction"] = str(e_lp_extract); module_logger.error(f"LabParser extract_labs_from_text error: {e_lp_extract}", exc_info=True)
    else: results.setdefault("lab_data", []); module_logger.warning(f"PIPELINE ID {content_hash}: LabParser not available.")

    # --- RecitationParser (extract_recitations_from_text) ---
    # ... (Keep RecitationParser logic as is from helpers_py_task_data_fix_v3) ...
    recitation_parser_instance = parsers.get('recitation_parser')
    if recitation_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 6 - Running RecitationParser.extract_recitations_from_text.")
        try:
            rec_text_segment = results.get("segmented_syllabus", {}).get("recitation_discussion_sections", "")
            if not isinstance(rec_text_segment, str) or not rec_text_segment.strip():
                rec_text_segment = syllabus_text_content 
                module_logger.info(f"Using full text segment for recitation extraction as 'recitation_discussion_sections' was empty.")
            if rec_text_segment:
                newly_extracted_recitations = recitation_parser_instance.extract_recitations_from_text(rec_text_segment, results.get("class_data", {}))
                if newly_extracted_recitations:
                    results.setdefault("recitation_data", []).extend(newly_extracted_recitations) 
                    results["metadata"]["Recitation Sessions"] = "Yes"
                    module_logger.info(f"RecitationParser.extract_recitations_from_text added {len(newly_extracted_recitations)} recitations. Total recitation_data: {len(results['recitation_data'])}.")
                else: module_logger.info("RecitationParser.extract_recitations_from_text returned no new recitations.")
            results["metadata"]["parser_stages_completed"].append("recitation_data_extraction")
        except Exception as e_rp_extract: results["metadata"]["error_recitation_parser_extraction"] = str(e_rp_extract); module_logger.error(f"RecitationParser extract_recitations_from_text error: {e_rp_extract}", exc_info=True)
    else: results.setdefault("recitation_data", []); module_logger.warning(f"PIPELINE ID {content_hash}: RecitationParser not available.")


    # --- ScheduleParser (generates base events, then enriches with LLM) ---
    schedule_parser_instance = parsers.get('schedule_parser')
    if schedule_parser_instance:
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 7 - Running ScheduleParser.process_schedule.")
        try: 
            results = schedule_parser_instance.process_schedule(results, unique_id=content_hash) 
            results["metadata"]["parser_stages_completed"].append("schedule_generation")
        except Exception as e_sch_proc: 
            results["metadata"]["error_schedule_parser_processing"] = str(e_sch_proc) 
            module_logger.error(f"ScheduleParser process_schedule error: {e_sch_proc}", exc_info=True)
    else:
        module_logger.warning(f"PIPELINE ID {content_hash}: ScheduleParser not available.")

    # --- Final processing stage for tasks (validation, linking) ---
    if assignment_parser_instance: # No need to check results.get("task_data") here, method handles empty
        module_logger.info(f"PIPELINE ID {content_hash}: Stage 8 - Running AssignmentParser.process_tasks_from_structured_data.")
        try:
            results = assignment_parser_instance.process_tasks_from_structured_data(results)
            results["metadata"]["parser_stages_completed"].append("assignment_processing_structured")
        except Exception as e_ap_struct:
            results["metadata"]["error_assignment_parser_structured"] = str(e_ap_struct)
            module_logger.error(f"AssignmentParser process_tasks_from_structured_data error: {e_ap_struct}", exc_info=True)

    # --- Final Metadata and Saving ---
    results["metadata"]["pipeline_end_time"] = datetime.now().isoformat()
    results["metadata"]["process_time_pipeline"] = round(time.monotonic() - pipeline_start_time, 3)
    results["metadata"]["extraction_complete"] = not any(key.startswith("error_") for key in results["metadata"])

    required_from_config = config.get("required_fields", [])
    final_missing_fields = [f for f in required_from_config if not results.get("class_data", {}).get(f, "").strip()]
    results["metadata"]["missing_fields"] = final_missing_fields
    
    module_logger.debug(f"PIPELINE ID {content_hash}: Final task_data before saving: {json.dumps(results.get('task_data', 'KEY_NOT_FOUND_OR_NONE'))}")
    module_logger.debug(f"PIPELINE ID {content_hash}: Final event_data before saving (first 3 items): {json.dumps(results.get('event_data', [])[:3], indent=2)}")


    parsed_results_output_dir = Path(config.get("directories", {}).get("parsed_syllabus_dir", "syllabus_data/parsed_results"))
    cfg_project_root_final_save = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
    if not parsed_results_output_dir.is_absolute():
        parsed_results_output_dir = cfg_project_root_final_save / parsed_results_output_dir
    parsed_results_output_dir.mkdir(parents=True, exist_ok=True) 

    final_json_path_to_save = parsed_results_output_dir / f"{content_hash}_results.json"

    if not write_json(final_json_path_to_save, results): 
        module_logger.error(f"CRITICAL: Failed to save final results to primary location: {final_json_path_to_save}")
    else:
        module_logger.info(f"Successfully saved final results to primary location: {final_json_path_to_save}")

    if transient_run_id: 
        transient_output_dir_str = config.get("directories", {}).get("output", "output") 
        cfg_project_root_transient = Path(config.get('project_root', Path(__file__).resolve().parent.parent))
        transient_output_dir = Path(transient_output_dir_str)
        if not transient_output_dir.is_absolute():
            transient_output_dir = cfg_project_root_transient / transient_output_dir
        transient_output_dir.mkdir(parents=True, exist_ok=True)
        
        transient_output_path = transient_output_dir / f"{transient_run_id}_pipeline_results.json" 
        if write_json(transient_output_path, results):
            module_logger.info(f"Also saved results to transient output: {transient_output_path} for run ID {transient_run_id}")
        else:
            module_logger.warning(f"Failed to save results to transient output: {transient_output_path}")

    module_logger.info(f"PIPELINE END: ID {content_hash}. Total Time: {results['metadata']['process_time_pipeline']:.2f}s. Success: {results['metadata']['extraction_complete']}. Missing required fields: {final_missing_fields}")
    return results

if __name__ == '__main__':
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s',
        handlers=[logging.StreamHandler(sys.stdout)]
    )
    module_logger.info("helpers.py executed directly. Running standalone tests/examples.")
    module_logger.info("\n--- helpers.py standalone tests finished (simplified) ---")

